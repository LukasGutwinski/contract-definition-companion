from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "demo" / "Contract-Definitions-Demo-SPA.docx"

INK = "171918"
NAVY = "020A2F"
TEAL = "0B5D78"
CYAN = "47D3F2"
MUTED = "626762"
LINE = "D8D8D0"
PAPER = "F5F7FB"
WHITE = "FFFFFF"
PALE_CYAN = "E1F7FB"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


DEFINITIONS = [
    ("Accounts", "the audited financial statements of the Company for the financial year ended on the Accounts Date."),
    ("Accounts Date", "31 December 2025."),
    ("Affiliate", "in relation to a person, any other person that directly or indirectly Controls, is Controlled by, or is under common Control with that person."),
    ("Agreement", "this fictional share purchase agreement, including its Schedules, as amended in writing by the Seller and the Buyer."),
    ("Applicable Law", "any statute, regulation, rule, order or binding requirement applicable to the relevant person, asset or circumstance."),
    ("Business", "the fictional business of designing and manufacturing precision components carried on by the Company as at Signing."),
    ("Business Day", "a day other than a Saturday, Sunday or public holiday in Vienna, Austria on which banks are generally open for business."),
    ("Buyer", "Example Buyer GmbH (fictional), as further identified on the first page of this Agreement."),
    ("Buyer Group", "the Buyer and each of its Affiliates from time to time, including the Company after Closing."),
    ("Claim", "a claim by the Buyer arising under or in connection with a Transaction Document."),
    ("Closing", "completion of the sale and purchase of the Shares in accordance with Clause 6."),
    ("Closing Date", "31 July 2026 or any other date agreed in writing by the Seller and the Buyer."),
    ("Closing Deliverables", "the documents and actions listed in Schedule 2."),
    ("Closing Payment", "the amount payable by the Buyer to the Seller at Closing under Clause 3.2."),
    ("Company", "Example Components GmbH (fictional), with fictitious company number FN 000002 c."),
    ("Company Group", "the Company and each Subsidiary listed in Schedule 1."),
    ("Confidential Information", "all non-public information relating to the Transaction, the Company Group, the Seller Group or the Buyer Group."),
    ("Consideration", "the aggregate value payable or provided by the Buyer for the Shares under this Agreement."),
    ("Control", "the direct or indirect power to direct the management or policies of a person, whether through ownership, voting rights, contract or otherwise."),
    ("Data Room", "the fictional electronic data room identified in Schedule 4, as indexed at 18:00 Vienna time on 20 July 2026."),
    ("Disclosure Letter", "the fictional disclosure letter delivered by the Seller to the Buyer immediately before Signing."),
    ("Encumbrance", "a mortgage, charge, pledge, lien, option, restriction, right of first refusal, security interest or other third-party right."),
    ("EUR", "the lawful currency of the member states of the European Union that have adopted the euro."),
    ("Fundamental Warranty", "a Seller Warranty set out in paragraphs 1 and 2 of Schedule 3."),
    ("Governmental Authority", "a court, regulator, ministry, municipality, competition authority, tax authority or other public body."),
    ("Intellectual Property", "patents, utility models, trade marks, designs, domain names, copyright, database rights, know-how and similar rights."),
    ("Leakage", "any transfer of value from the Company Group to the Seller Group between the Locked Box Date and Closing, other than Permitted Leakage."),
    ("Locked Box Accounts", "the unaudited statement of financial position of the Company Group as at the Locked Box Date included in folder 2.1 of the Data Room."),
    ("Locked Box Date", "31 March 2026."),
    ("Loss", "a direct loss, liability, damage, cost or expense reasonably incurred, excluding punitive damages and remote or indirect loss."),
    ("Material Adverse Effect", "an event or circumstance that has, or is reasonably likely to have, a material adverse effect on the Business, assets or financial condition of the Company Group taken as a whole, excluding general market changes and matters fairly disclosed in the Data Room."),
    ("Notary", "the Austrian civil law notary selected jointly by the Seller and the Buyer for the notarial implementation of the Transaction."),
    ("Permitted Leakage", "a payment or benefit expressly listed in Schedule 5."),
    ("Purchase Price", "EUR 12,500,000, subject only to the adjustments expressly set out in this Agreement."),
    ("Records", "the books, contracts, personnel files, accounting records and other records relating primarily to the Company Group."),
    ("Regulatory Approval", "the fictional merger-control clearance described in Clause 5.1."),
    ("Seller", "Example Seller GmbH (fictional), as further identified on the first page of this Agreement."),
    ("Seller Group", "the Seller and each of its Affiliates from time to time, excluding the Company Group after Closing."),
    ("Seller Warranties", "the warranties given by the Seller under Clause 8 and set out in Schedule 3."),
    ("Shares", "the single share representing 100 percent of the registered share capital of the Company, as described in Schedule 1."),
    ("Signing", "the execution of this Agreement by the Seller and the Buyer for demonstration purposes only."),
    ("Subsidiary", "an entity Controlled directly or indirectly by the Company."),
    ("Tax", "any tax, levy, duty, social security contribution or similar governmental charge, together with related interest and penalties."),
    ("Tax Authority", "a Governmental Authority responsible for the assessment, administration or collection of Tax."),
    ("Tax Claim", "a Claim under the Tax Covenant or for breach of a Seller Warranty relating to Tax."),
    ("Tax Covenant", "the fictional covenant in Clause 11 concerning Tax attributable to periods ending on or before Closing."),
    ("Transaction", "the proposed sale and purchase of the Shares contemplated by this Agreement."),
    ("Transaction Documents", "this Agreement, the Disclosure Letter and each document identified as such in Schedule 2."),
    ("Warranty", "a Seller Warranty or a warranty given by the Buyer under Clause 9."),
    ("Warranty Claim", "a Claim for breach of a Seller Warranty."),
    ("Working Hours", "the period from 09:00 to 17:30 on a Business Day in Vienna, Austria."),
]


CLAUSES = [
    (
        "SALE AND PURCHASE",
        [
            (
                "Sale of the Shares",
                [
                    "Subject to the terms of this Agreement, the Seller shall sell and the Buyer shall purchase the Shares at Closing free from every Encumbrance and together with all rights attaching to the Shares at Closing.",
                    "The Buyer is not obliged to complete the purchase of any part of the Shares unless the purchase of all Shares is completed simultaneously.",
                ],
            ),
            (
                "Title and economic benefit",
                [
                    "Legal and economic title to the Shares shall pass to the Buyer at Closing upon completion of the Closing Deliverables before the Notary and payment of the Closing Payment.",
                ],
            ),
        ],
    ),
    (
        "PURCHASE PRICE AND PAYMENT",
        [
            (
                "Purchase Price",
                [
                    "The Purchase Price for the Shares is EUR 12,500,000. The Purchase Price has been agreed by reference to the Locked Box Accounts and assumes that no Leakage occurs.",
                ],
            ),
            (
                "Closing Payment",
                [
                    "At Closing, the Buyer shall pay the Closing Payment to the bank account notified by the Seller at least two Business Days before the Closing Date.",
                    "Payment of the Closing Payment in accordance with this Clause discharges the Buyer from its corresponding payment obligation. No amount may be withheld or set off except as required by Applicable Law.",
                ],
            ),
            (
                "Leakage",
                [
                    "The Seller undertakes that no Leakage has occurred since the Locked Box Date and that no Leakage will occur before Closing, other than Permitted Leakage.",
                    "If Leakage occurs, the Seller shall pay the Buyer an amount equal to that Leakage within ten Business Days after receiving reasonable supporting details from the Buyer.",
                ],
            ),
        ],
    ),
    (
        "CONDUCT BEFORE CLOSING",
        [
            (
                "Ordinary course",
                [
                    "Between Signing and Closing, the Seller shall procure that the Company Group carries on the Business in the ordinary course and preserves its material relationships with employees, customers and suppliers.",
                    "The Seller shall promptly notify the Buyer of any event that constitutes, or could reasonably be expected to constitute, a Material Adverse Effect.",
                ],
            ),
            (
                "Restricted actions",
                [
                    "Without the Buyer's prior written consent, the Seller shall procure that no member of the Company Group declares a distribution, incurs material financial indebtedness outside the ordinary course, disposes of a material asset, grants an Encumbrance or varies the rights attaching to the Shares.",
                ],
            ),
            (
                "Access",
                [
                    "During Working Hours and on reasonable notice, the Seller shall provide the Buyer and its advisers with reasonable access to senior management and Records, provided that access does not materially disrupt the Business or breach Applicable Law.",
                ],
            ),
        ],
    ),
    (
        "CONDITIONS",
        [
            (
                "Conditions to Closing",
                [
                    "Closing is conditional on the Regulatory Approval having been obtained or deemed obtained and on the Notary confirming that the agreed notarial implementation steps can be completed on the Closing Date.",
                    "Each party shall use reasonable endeavours to satisfy the conditions as soon as reasonably practicable and shall keep the other party informed of material communications with a Governmental Authority.",
                ],
            ),
            (
                "Failure of conditions",
                [
                    "If the conditions are not satisfied or waived by 30 September 2026, either party may terminate this Agreement by written notice, except where that party's material breach caused the failure.",
                ],
            ),
        ],
    ),
    (
        "CLOSING",
        [
            (
                "Time and place",
                [
                    "Closing shall take place on the Closing Date at the office of the Notary in Vienna, or remotely to the extent permitted by Applicable Law and accepted by the Notary.",
                ],
            ),
            (
                "Closing actions",
                [
                    "At Closing, the Seller and the Buyer shall perform the Closing Deliverables in the sequence set out in Schedule 2. The Closing Deliverables are interdependent and, so far as legally possible, are deemed to occur simultaneously.",
                    "If a material Closing Deliverable is not completed, the non-defaulting party may defer Closing for up to five Business Days, proceed so far as practicable, or terminate this Agreement where the failure is not remedied within that period.",
                ],
            ),
        ],
    ),
    (
        "POST-CLOSING OBLIGATIONS",
        [
            (
                "Records and assistance",
                [
                    "For seven years after Closing, each party shall retain Records in its possession that the other party may reasonably require for Tax, accounting or regulatory purposes and shall provide reasonable access during Working Hours.",
                ],
            ),
            (
                "Company name and branding",
                [
                    "Within 60 days after Closing, the Buyer shall procure that the Company Group ceases using any trade name or logo belonging exclusively to the Seller Group, except where continued use is expressly permitted by a Transaction Document.",
                ],
            ),
        ],
    ),
    (
        "SELLER WARRANTIES",
        [
            (
                "Warranties at Signing",
                [
                    "The Seller warrants to the Buyer that each Seller Warranty is true and accurate at Signing, subject to matters fairly disclosed in the Disclosure Letter.",
                    "Each Seller Warranty is separate and is not limited by reference to another Seller Warranty, except where this Agreement expressly provides otherwise.",
                ],
            ),
            (
                "Warranties at Closing",
                [
                    "The Fundamental Warranties are deemed repeated at Closing by reference to the facts and circumstances then existing. The Seller shall notify the Buyer before Closing of any matter that would make a Fundamental Warranty untrue or misleading.",
                ],
            ),
            (
                "No implied warranties",
                [
                    "Except for the Seller Warranties and the Tax Covenant, the Buyer acknowledges that it has not relied on any representation concerning the Transaction, the Business or the Company Group.",
                ],
            ),
        ],
    ),
    (
        "BUYER WARRANTIES",
        [
            (
                "Authority and funds",
                [
                    "The Buyer warrants that it has authority to enter into and perform the Transaction Documents and has sufficient immediately available funds to pay the Closing Payment.",
                ],
            ),
            (
                "No conflict",
                [
                    "The Buyer warrants that its entry into and performance of the Transaction Documents does not breach its constitutional documents, Applicable Law or a material agreement binding on the Buyer.",
                ],
            ),
        ],
    ),
    (
        "LIMITATIONS ON CLAIMS",
        [
            (
                "Notification",
                [
                    "The Buyer shall notify the Seller of a Warranty Claim with reasonable details of the nature of the Claim and, so far as reasonably practicable, the estimated Loss.",
                ],
            ),
            (
                "Time limits",
                [
                    "A Warranty Claim other than a Tax Claim or a Claim for breach of a Fundamental Warranty must be notified within 18 months after the Closing Date. A Tax Claim must be notified within six years after the Closing Date.",
                ],
            ),
            (
                "Financial thresholds",
                [
                    "The Seller is not liable for an individual Warranty Claim below EUR 10,000 unless related Claims together exceed that amount. The Seller is not liable for Warranty Claims until the aggregate qualifying Loss exceeds EUR 125,000, after which the Seller is liable only for the excess.",
                    "The Seller's aggregate liability for Warranty Claims is limited to 20 percent of the Purchase Price. Liability for breach of a Fundamental Warranty, fraud or deliberate concealment is not limited by that cap to the extent prohibited by Applicable Law.",
                ],
            ),
            (
                "No double recovery",
                [
                    "The Buyer may not recover more than once for the same Loss and shall take reasonable steps to mitigate any Loss giving rise to a Claim.",
                ],
            ),
        ],
    ),
    (
        "TAX COVENANT",
        [
            (
                "Seller covenant",
                [
                    "Subject to this Agreement, the Seller shall pay the Buyer an amount equal to Tax of the Company Group attributable to an event occurring or income earned on or before Closing, except to the extent that the Tax was specifically provided for in the Locked Box Accounts.",
                ],
            ),
            (
                "Tax conduct",
                [
                    "The Buyer shall keep the Seller reasonably informed of a Tax Claim and shall not settle a material Tax Claim relating to a pre-Closing period without consulting the Seller, unless immediate action is required by Applicable Law or the Tax Authority.",
                ],
            ),
        ],
    ),
    (
        "RESTRICTIVE COVENANTS",
        [
            (
                "Non-compete",
                [
                    "For 18 months after Closing, the Seller shall not knowingly carry on a business in Austria that competes directly with the Business as carried on at Closing, except through a passive holding of less than five percent in a listed company.",
                ],
            ),
            (
                "Non-solicitation",
                [
                    "For 18 months after Closing, the Seller shall not knowingly solicit for employment a senior employee of the Company Group or solicit a material customer to cease doing business with the Company Group.",
                ],
            ),
        ],
    ),
    (
        "CONFIDENTIALITY AND ANNOUNCEMENTS",
        [
            (
                "Confidentiality",
                [
                    "Each party shall keep Confidential Information confidential and may disclose it only to its Affiliates, professional advisers, financiers or a Governmental Authority where disclosure is required for the Transaction or by Applicable Law.",
                ],
            ),
            (
                "Announcements",
                [
                    "No public announcement concerning the Transaction may be made before Closing without the other party's prior written approval, except where required by Applicable Law. After Closing, the Buyer controls announcements concerning the Company Group.",
                ],
            ),
        ],
    ),
    (
        "DATA PROTECTION",
        [
            (
                "Transaction data",
                [
                    "Each party shall process personal data received in connection with the Transaction in accordance with Applicable Law and shall apply appropriate technical and organisational measures to protect that data.",
                ],
            ),
            (
                "Data Room",
                [
                    "The Buyer shall use personal data in the Data Room only for evaluating and implementing the Transaction unless another lawful basis applies.",
                ],
            ),
        ],
    ),
    (
        "NOTICES",
        [
            (
                "Form of notice",
                [
                    "A notice under a Transaction Document must be in writing and delivered by hand, reputable courier or email to the address notified by the relevant party.",
                ],
            ),
            (
                "Deemed receipt",
                [
                    "A notice delivered by hand is deemed received when left at the notified address. A courier notice is deemed received at 10:00 on the second Business Day after dispatch. An email is deemed received when sent, unless sent outside Working Hours, in which case it is deemed received at 09:00 on the next Business Day.",
                ],
            ),
        ],
    ),
    (
        "COSTS, ASSIGNMENT AND THIRD-PARTY RIGHTS",
        [
            (
                "Costs",
                [
                    "Except as otherwise stated in a Transaction Document, each party shall bear its own costs relating to the Transaction. The Buyer shall bear the fees of the Notary for implementing the transfer of the Shares.",
                ],
            ),
            (
                "Assignment",
                [
                    "Neither party may assign a right under this Agreement without the other party's prior written consent, except that the Buyer may assign to an Affiliate while remaining liable for performance.",
                ],
            ),
            (
                "Third-party rights",
                [
                    "A person that is not a party to this Agreement has no right to enforce a term of this Agreement, except where a Transaction Document expressly provides otherwise.",
                ],
            ),
        ],
    ),
    (
        "ENTIRE AGREEMENT AND REMEDIES",
        [
            (
                "Entire agreement",
                [
                    "The Transaction Documents constitute the entire agreement between the parties concerning the Transaction and supersede prior drafts, discussions and understandings concerning their subject matter.",
                ],
            ),
            (
                "Remedies",
                [
                    "The rights and remedies in this Agreement are cumulative unless expressly stated otherwise. A waiver is effective only if made in writing and applies only to the specific circumstances for which it is given.",
                ],
            ),
        ],
    ),
    (
        "MISCELLANEOUS",
        [
            (
                "Further assurance",
                [
                    "After Closing, each party shall execute documents and take actions reasonably required to give full effect to the Transaction and the transfer of the Shares.",
                ],
            ),
            (
                "Severability",
                [
                    "If a provision of this Agreement is invalid or unenforceable, it shall be modified to the minimum extent necessary and the remaining provisions shall continue in effect.",
                ],
            ),
            (
                "Counterparts",
                [
                    "This Agreement may be executed in counterparts for demonstration purposes, but the parties acknowledge that the actual transfer of shares in an Austrian limited liability company may require implementation in the form required by Applicable Law before the Notary.",
                ],
            ),
        ],
    ),
    (
        "GOVERNING LAW AND JURISDICTION",
        [
            (
                "Governing law",
                [
                    "This Agreement and any non-contractual obligations arising out of or in connection with it are governed by the laws of Austria.",
                ],
            ),
            (
                "Jurisdiction",
                [
                    "Subject to mandatory Applicable Law, the courts competent for Vienna, Inner City, have exclusive jurisdiction to settle a dispute arising out of or in connection with this Agreement.",
                    "This Clause is included solely to make the fictional document resemble a transaction document. It is not a recommendation for any real transaction.",
                ],
            ),
        ],
    ),
]


def set_run_font(run, name="Times New Roman", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=INK, bold=False):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGINS_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        tag = qn(f"w:{side}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run_node = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "16")
    props.append(color)
    props.append(size)
    run_node.append(props)
    text = OxmlElement("w:t")
    text.text = "1"
    run_node.append(text)
    field.append(run_node)
    paragraph._p.append(field)


def next_numbering_id(numbering, tag, attr):
    values = []
    for element in numbering.findall(qn(tag)):
        value = element.get(qn(attr))
        if value is not None:
            values.append(int(value))
    return max(values, default=-1) + 1


def create_clause_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    formats = [
        (0, "%1.", 0, 360, 360),
        (1, "%1.%2", 360, 720, 360),
        (2, "%1.%2.%3", 720, 1080, 360),
    ]
    for level, text_value, left, text_indent, hanging in formats:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_indent))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
        abstract.append(lvl)

    numbering.append(abstract)
    num_id = next_numbering_id(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def create_recital_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "upperLetter")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, suff, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num_id = next_numbering_id(numbering, "w:num", "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_body(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    return paragraph


def add_clause_heading(doc, title, clause_number, subclause_number=None):
    level = 0 if subclause_number is None else 1
    style_name = f"Heading {level + 1}"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.keep_with_next = True
    number_text = (
        f"{clause_number}."
        if subclause_number is None
        else f"{clause_number}.{subclause_number}"
    )
    paragraph.add_run(f"{number_text} {title}")
    return paragraph


def add_definition(doc, term, meaning):
    paragraph = doc.add_paragraph(style="Definition")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.keep_together = True
    term_run = paragraph.add_run(f'"{term}"')
    term_run.bold = True
    paragraph.add_run(" means")
    paragraph.add_run().add_break()
    paragraph.add_run(meaning)
    return paragraph


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.style = "Table Header"
        paragraph.add_run(header)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = "Table Text"
            paragraph.add_run(value)
    return table


def add_notice_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_CYAN)
    paragraph = cell.paragraphs[0]
    paragraph.style = "Demo Notice"
    paragraph.add_run(text)
    return table


def add_page_break_with_spacer(doc):
    doc.add_page_break()
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("\u00a0")
    set_run_font(run, size=1, color=WHITE)


def add_schedule_title(doc, title, subtitle=None):
    title_paragraph = doc.add_paragraph(style="Schedule Heading")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.left_indent = Inches(0)
    title_paragraph.paragraph_format.right_indent = Inches(0)
    title_paragraph.paragraph_format.first_line_indent = Inches(0)
    title_paragraph.add_run(title)
    if subtitle:
        subtitle_paragraph = doc.add_paragraph(style="Schedule Subheading")
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitle_paragraph.paragraph_format.left_indent = Inches(0)
        subtitle_paragraph.paragraph_format.right_indent = Inches(0)
        subtitle_paragraph.paragraph_format.first_line_indent = Inches(0)
        subtitle_paragraph.paragraph_format.keep_with_next = True
        subtitle_paragraph.add_run(subtitle)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Times New Roman", 10.5, INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 12, TEAL, 8, 4),
    ]:
        style = styles[name]
        set_style_font(style, "Times New Roman", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    custom_styles = {
        "Cover Kicker": ("Times New Roman", 10, CYAN, True, 0, 14, WD_ALIGN_PARAGRAPH.CENTER),
        "Cover Title": ("Times New Roman", 28, NAVY, True, 0, 8, WD_ALIGN_PARAGRAPH.CENTER),
        "Cover Subtitle": ("Times New Roman", 13, MUTED, False, 0, 24, WD_ALIGN_PARAGRAPH.CENTER),
        "Cover Party": ("Times New Roman", 11, INK, False, 0, 4, WD_ALIGN_PARAGRAPH.CENTER),
        "Unnumbered Heading": ("Times New Roman", 16, NAVY, True, 16, 8, WD_ALIGN_PARAGRAPH.LEFT),
        "Definition": ("Times New Roman", 10.5, INK, False, 0, 4, WD_ALIGN_PARAGRAPH.LEFT),
        "Schedule Heading": ("Times New Roman", 19, NAVY, True, 0, 16, WD_ALIGN_PARAGRAPH.CENTER),
        "Schedule Subheading": ("Times New Roman", 13, TEAL, True, 12, 6, WD_ALIGN_PARAGRAPH.LEFT),
        "Demo Notice": ("Times New Roman", 10, NAVY, True, 0, 0, WD_ALIGN_PARAGRAPH.LEFT),
        "Table Text": ("Times New Roman", 9.5, INK, False, 0, 0, WD_ALIGN_PARAGRAPH.LEFT),
        "Table Header": ("Times New Roman", 9.5, WHITE, True, 0, 0, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for name, config in custom_styles.items():
        style = styles.add_style(name, 1)
        font, size, color, bold, before, after, alignment = config
        set_style_font(style, font, size, color, bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.alignment = alignment
        style.paragraph_format.widow_control = True
        if name == "Schedule Subheading":
            style.paragraph_format.keep_with_next = True

    definition = styles["Definition"]
    definition.paragraph_format.left_indent = Inches(0)
    definition.paragraph_format.first_line_indent = Inches(0)
    definition.paragraph_format.keep_together = True


def populate_header(header):
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    spacer = paragraph.add_run("\u00a0")
    set_run_font(spacer, size=1, color=WHITE)


def populate_footer(footer):
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(paragraph)


def configure_section(doc, section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Define both page variants explicitly. This avoids office-suite-specific
    # interpretation of a single default header on alternating pages.
    doc.settings.odd_and_even_pages_header_footer = True
    populate_header(section.header)
    populate_header(section.even_page_header)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def build_document():
    doc = Document()
    section = doc.sections[0]
    configure_section(doc, section)
    configure_styles(doc)

    props = doc.core_properties
    props.title = "Contract Definitions Demo Share Purchase Agreement"
    props.subject = "Fictional SPA for software demonstration and Microsoft Marketplace testing"
    props.author = "GUT Ventures GmbH"
    props.keywords = "demo, share purchase agreement, defined terms, Word add-in"
    props.comments = "Fictional document. Not legal advice and not for signature."

    for _ in range(4):
        doc.add_paragraph()

    doc.add_paragraph("DEMO DOCUMENT", style="Cover Kicker")
    doc.add_paragraph("SHARE PURCHASE AGREEMENT", style="Cover Title")
    doc.add_paragraph(
        "Sale and purchase of 100% of the shares in Example Components GmbH (fictional)",
        style="Cover Subtitle",
    )

    doc.add_paragraph("between", style="Cover Party")
    seller = doc.add_paragraph(style="Cover Party")
    seller_run = seller.add_run("EXAMPLE SELLER GMBH")
    seller_run.bold = True
    seller.add_run(" (fictional) as Seller")
    doc.add_paragraph("and", style="Cover Party")
    buyer = doc.add_paragraph(style="Cover Party")
    buyer_run = buyer.add_run("EXAMPLE BUYER GMBH")
    buyer_run.bold = True
    buyer.add_run(" (fictional) as Buyer")

    date_paragraph = doc.add_paragraph(style="Cover Party")
    date_paragraph.paragraph_format.space_before = Pt(18)
    date_run = date_paragraph.add_run("24 July 2026")
    date_run.bold = True

    doc.add_paragraph()
    add_notice_box(
        doc,
        "IMPORTANT: This document is entirely fictional and was created solely to demonstrate and test Contract Definitions for Word. It is not legal advice, is not intended for execution and must not be used as a transaction template.",
    )
    doc.add_paragraph()
    demo_note = doc.add_paragraph()
    demo_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = demo_note.add_run("No real persons, companies, registration numbers, addresses or commercial terms are used.")
    set_run_font(note_run, size=9, color=MUTED, italic=True)
    add_page_break_with_spacer(doc)

    doc.add_paragraph("PARTIES", style="Unnumbered Heading")
    add_body(
        doc,
        "Example Seller GmbH (fictional), a fictional company described for demonstration purposes as incorporated under the laws of Austria with fictitious company number FN 000000 a and fictitious registered office at Demo Street 1, 1010 Vienna, Austria (the \"Seller\").",
    )
    add_body(
        doc,
        "Example Buyer GmbH (fictional), a fictional company described for demonstration purposes as incorporated under the laws of Austria with fictitious company number FN 000001 b and fictitious registered office at Sample Avenue 2, 1010 Vienna, Austria (the \"Buyer\").",
    )

    doc.add_paragraph("RECITALS", style="Unnumbered Heading")
    recitals = [
        "The Seller is the legal and beneficial owner of the Shares.",
        "The Company carries on the Business and owns the Subsidiary identified in Schedule 1.",
        "The Buyer wishes to acquire, and the Seller wishes to sell, the Shares on the terms of this Agreement.",
        "The parties intend the Transaction to be implemented before the Notary in the form required by Applicable Law.",
    ]
    for recital_index, recital in enumerate(recitals):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.add_run(f"{chr(ord('A') + recital_index)}. {recital}")

    add_clause_heading(doc, "DEFINITIONS AND INTERPRETATION", 1)
    add_clause_heading(doc, "Definitions", 1, 1)
    add_body(
        doc,
        "In this Agreement, the following capitalised words and expressions have the meanings set out below.",
    )
    for definition_index, (term, meaning) in enumerate(DEFINITIONS, start=1):
        add_definition(doc, term, meaning)
        if definition_index in {6, 20, 34}:
            doc.add_page_break()

    add_clause_heading(doc, "Interpretation", 1, 2)
    add_body(
        doc,
        "A reference to a Clause or Schedule is to a clause of or schedule to this Agreement. Headings do not affect interpretation. The words \"including\" and \"includes\" do not limit the preceding words. A reference to writing includes email, except for a notice under Clause 15.",
    )
    add_body(
        doc,
        "An obligation on the Seller to procure an act by the Company applies only before Closing. After Closing, an obligation on the Buyer to procure an act by the Company applies for so long as the Buyer Controls the Company.",
    )

    for clause_number, (clause_title, subclauses) in enumerate(CLAUSES, start=2):
        if clause_number in {14, 17, 19}:
            add_page_break_with_spacer(doc)
        add_clause_heading(doc, clause_title, clause_number)
        for subclause_number, (subheading, paragraphs) in enumerate(subclauses, start=1):
            add_clause_heading(doc, subheading, clause_number, subclause_number)
            for text in paragraphs:
                add_body(doc, text)

    doc.add_paragraph()
    add_notice_box(
        doc,
        "NO EXECUTION: The signature page is intentionally omitted. This fictional Demo Agreement cannot be signed or used to transfer any shares.",
    )

    add_page_break_with_spacer(doc)
    add_schedule_title(doc, "SCHEDULE 1", "THE COMPANY, SUBSIDIARY AND SHARES")
    add_table(
        doc,
        ["Item", "Fictional details"],
        [
            ("Company", "Example Components GmbH (fictional)"),
            ("Fictitious company number", "FN 000002 c"),
            ("Fictitious registered office", "Prototype Road 3, 4020 Linz, Austria"),
            ("Registered share capital", "EUR 35,000"),
            ("Shares sold", "One share representing 100 percent of the registered share capital"),
            ("Seller", "Example Seller GmbH (fictional)"),
            ("Subsidiary", "Example Components Services s.r.o. (fictional)"),
            ("Business", "Design and manufacture of fictional precision components"),
        ],
        [2400, 6960],
    )

    add_page_break_with_spacer(doc)
    add_schedule_title(doc, "SCHEDULE 2", "CLOSING DELIVERABLES")
    add_table(
        doc,
        ["Step", "Responsible party", "Closing Deliverable"],
        [
            ("1", "Seller", "Evidence reasonably satisfactory to the Notary of the Seller's authority to complete the Transaction."),
            ("2", "Seller and Buyer", "Notarial implementation document in the form agreed with the Notary."),
            ("3", "Seller", "Written resignations of the fictional managing directors specified by the Buyer, effective at Closing."),
            ("4", "Buyer", "Evidence of irrevocable payment instructions for the Closing Payment."),
            ("5", "Seller", "The Disclosure Letter and an electronic copy of the Data Room index."),
            ("6", "Company", "Updated shareholder records and the filings required under Applicable Law following Closing."),
            ("7", "Seller and Buyer", "A joint Closing memorandum recording completion of the Closing Deliverables."),
        ],
        [760, 1900, 6700],
    )

    add_page_break_with_spacer(doc)
    add_schedule_title(doc, "SCHEDULE 3", "SELLER WARRANTIES")
    warranty_sections = [
        (
            "1. CAPACITY AND TITLE",
            [
                "The Seller has authority to enter into and perform the Transaction Documents.",
                "The Seller is the sole legal and beneficial owner of the Shares. The Shares are fully paid and free from every Encumbrance.",
            ],
        ),
        (
            "2. THE COMPANY GROUP",
            [
                "The information in Schedule 1 is accurate in all material respects.",
                "No member of the Company Group is insolvent or subject to a pending resolution for liquidation.",
            ],
        ),
        (
            "3. ACCOUNTS AND RECORDS",
            [
                "The Accounts have been prepared consistently with the accounting policies stated in the Accounts and give a true and fair view in all material respects for demonstration purposes.",
                "The Records are maintained in all material respects in accordance with Applicable Law and are in the possession or control of the Company Group.",
            ],
        ),
        (
            "4. MATERIAL CONTRACTS",
            [
                "The Data Room contains copies of each material customer and supplier contract of the Company Group.",
                "No member of the Company Group has received written notice alleging a material unremedied breach of a material contract.",
            ],
        ),
        (
            "5. EMPLOYEES",
            [
                "The Data Room contains an anonymised list of employees showing role, start date and fixed remuneration.",
                "No material collective dispute is current or, so far as the Seller is aware, threatened in writing.",
            ],
        ),
        (
            "6. INTELLECTUAL PROPERTY AND IT",
            [
                "The Company Group owns or is licensed to use the Intellectual Property material to the Business.",
                "During the 12 months before Signing, the Company Group has not suffered a material, unremedied outage or cybersecurity incident affecting the Business.",
            ],
        ),
        (
            "7. COMPLIANCE AND LITIGATION",
            [
                "The Company Group has not received written notice from a Governmental Authority alleging a material breach of Applicable Law that remains unresolved.",
                "No material litigation involving a member of the Company Group is pending or, so far as the Seller is aware, threatened in writing.",
            ],
        ),
        (
            "8. TAX",
            [
                "The Company Group has filed all material Tax returns required to be filed and has paid all material Tax due, subject to the matters disclosed in the Disclosure Letter.",
                "The Company Group is not engaged in a material dispute with a Tax Authority.",
            ],
        ),
    ]
    for warranty_index, (heading, paragraphs) in enumerate(warranty_sections, start=1):
        if warranty_index == 5:
            add_page_break_with_spacer(doc)
        doc.add_paragraph(heading, style="Schedule Subheading")
        for text in paragraphs:
            add_body(doc, text)

    add_page_break_with_spacer(doc)
    add_schedule_title(doc, "SCHEDULE 4", "DATA ROOM INDEX")
    add_table(
        doc,
        ["Folder", "Category", "Illustrative contents"],
        [
            ("1.0", "Corporate", "Fictional articles, shareholder records and group chart"),
            ("2.0", "Financial", "Accounts, Locked Box Accounts and management reporting"),
            ("3.0", "Commercial", "Anonymised customer and supplier contract summaries"),
            ("4.0", "Employees", "Anonymised employee schedule and template employment agreements"),
            ("5.0", "Intellectual Property", "Trade mark list, software inventory and licence summaries"),
            ("6.0", "Real estate", "Fictional lease summaries for the Linz premises"),
            ("7.0", "Compliance", "Policies, permits and correspondence summaries"),
            ("8.0", "Tax", "Tax return summaries and fictional Tax Authority correspondence"),
        ],
        [1100, 2400, 5860],
    )
    add_body(
        doc,
        "The Data Room and this index are fictional. They are included only to create realistic occurrences of defined terms for Contract Definitions for Word.",
    )

    add_page_break_with_spacer(doc)
    add_schedule_title(doc, "SCHEDULE 5", "PERMITTED LEAKAGE")
    add_table(
        doc,
        ["Item", "Maximum amount"],
        [
            ("Salary and contractual benefits paid to Seller-nominated managing directors in the ordinary course", "EUR 45,000"),
            ("Advisory fees payable by the Company to the Seller under the disclosed transition services arrangement", "EUR 25,000"),
            ("Reimbursement of documented third-party costs incurred for the Transaction", "EUR 15,000"),
        ],
        [7000, 2360],
    )
    add_body(
        doc,
        "Any payment that exceeds the corresponding maximum amount is Leakage to the extent of the excess and is not Permitted Leakage.",
    )

    add_page_break_with_spacer(doc)
    add_schedule_title(doc, "DEMO USE NOTE")
    add_notice_box(
        doc,
        "This fictional SPA was authored for Contract Definitions for Word. It deliberately contains a substantial Definitions section and repeated defined terms throughout the clauses and Schedules so that definition detection, highlighting and occurrence navigation can be demonstrated.",
    )
    add_body(
        doc,
        "The document may be used in screenshots, screen recordings and Microsoft Marketplace certification testing for Contract Definitions for Word. It contains no real transaction data and is not intended to reflect a complete or legally sufficient Austrian share transfer instrument.",
    )
    add_body(
        doc,
        "For a real transaction, parties should obtain advice from qualified legal and tax advisers and comply with all formal requirements under Applicable Law.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
